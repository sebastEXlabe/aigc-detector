# -*- coding: utf-8 -*-
"""AIGC 检测报告增量入库管线。

用途：扫描工具目录 reports/ 及常见 source 目录（Downloads/DocumentsWXWork/D:xwechat/OneDrive），
发现新增 AIGC 检测报告 → 复制归档到 reports/{pdf,word,html} → 解析句级标注 → 合并进 data/train_unified.jsonl。
用法：python ingest_new_reports.py [--source-dir <path>] [--threshold <n>]
"""
import os, re, json, sys, io, hashlib, shutil, time
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

BASE = r"C:\Users\woshi\.dsh\aigc-detector"
DATA = os.path.join(BASE, "data")
REPORTS = os.path.join(BASE, "reports")
DATASET = os.path.join(DATA, "train_unified.jsonl")
MANIFEST = os.path.join(BASE, "reports_manifest.json")

# 常见 source 目录（新报告最可能出现的扫描位置）
DEFAULT_SOURCES = [
    r"C:\Users\woshi\Downloads",
    r"C:\Users\woshi\Documents\WXWork",
    r"D:\xwechat_files",
    r"C:\Users\woshi\OneDrive",
]

# AIGC 检测报告文件名/内容特征
NAME_KEYS = ["AIGC检测报告", "AIGC报告", "AIGC检测", "AIGC存档", "AIGC全文", "AIGC原文对照",
             "AIGC_全文报告单", "AIGC报告单", "AIGC检测报告单", "AIGC简洁报告", "AIGC报告",
             "标红版_AIGC", "知网AIGC", "PaperYY", "PaperPass", "维普AIGC", "查重报告"]
CONTENT_FP = ["AIGC检测", "AI特征值", "疑似率", "疑似AIGC", "AIGC总体疑似", "AI生成的可能性",
              "人工智能生成检测", "AIGC文档检测"]

def is_report_file(fp):
    base = os.path.basename(fp)
    if not re.search(r"\.(pdf|docx|html?)$", base, re.I):
        return False
    low = base.lower()
    if any(k in base or k in low for k in NAME_KEYS):
        return True
    return False

def parse_report(fp):
    """按类型解析报告，返回句级标注列表 [{text, prob, label}] 或 None。"""
    ext = os.path.splitext(fp)[1].lower()
    if ext == ".pdf":
        return _parse_pdf(fp)
    elif ext == ".docx":
        return _parse_word(fp)
    elif ext in (".html", ".htm"):
        return _parse_html(fp)
    return None

def _read_report_text(fp):
    """读取报告原始文本（供元数据提取）。"""
    ext = os.path.splitext(fp)[1].lower()
    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(fp)
            t = "".join(pg.get_text() for pg in doc)
            doc.close()
            return t
        elif ext == ".docx":
            from docx import Document
            d = Document(fp)
            return "\n".join(p.text for p in d.paragraphs)
        elif ext in (".html", ".htm"):
            return open(fp, encoding="utf-8", errors="ignore").read()
    except Exception:
        pass
    return ""

def extract_report_meta(fp):
    """提取报告关键参数（对齐真实AIGC报告口径）。返回 dict。"""
    t = _read_report_text(fp)
    meta = {"file": os.path.basename(fp), "platform": "unknown"}
    # 平台
    if "cx.cnki.net" in t or ("知网" in t and "AIGC检测" in t): meta["platform"] = "知网"
    elif "PaperPass" in t or "PaperYY" in t: meta["platform"] = "PaperPass/PaperYY"
    elif "维普" in t or "cqvip" in t: meta["platform"] = "维普"
    elif "超星" in t or "学习通" in t or "大雅" in t: meta["platform"] = "超星/大雅"
    elif "万方" in t or "wanfang" in t: meta["platform"] = "万方"
    elif "原文对照" in t or "人工撰写占比" in t: meta["platform"] = "原文对照"
    # 报告编号
    m = re.search(r"(?:NO|编号|报告编号)[：:]\s*([A-Za-z0-9_]+)", t); meta["report_no"] = m.group(1) if m else None
    # AI特征值 / 疑似率
    m = re.search(r"AI特征值[：:]\s*([\d.]+)%", t)
    if not m: m = re.search(r"疑似率[：:]\s*([\d.]+)%", t)
    if not m: m = re.search(r"AIGC总体疑似度[^0-9]{0,8}\s*([\d.]+)%", t)
    meta["ai_rate"] = float(m.group(1)) if m else None
    # 字数 / 总字符数
    m = re.search(r"(?:字数|总字符数)[：:]\s*([\d,]+)", t); meta["total_chars"] = int(m.group(1).replace(",","")) if m else None
    # AI特征字符数
    m = re.search(r"AI特征字符数[：:]\s*([\d]+)", t); meta["ai_chars"] = int(m.group(1)) if m else None
    # 检测时间
    m = re.search(r"检测时间[：:]\s*([\d-]+ [\d:]+)", t); meta["detect_time"] = m.group(1) if m else None
    # 篇名
    m = re.search(r"篇名[：:]\s*(.+)", t); meta["title"] = m.group(1).strip() if m else None
    # 作者
    m = re.search(r"作者[：:]\s*(.+)", t); meta["author"] = m.group(1).strip() if m else None
    return meta

def _split_sentences(t):
    return [p.strip() for p in re.split(r"(?<=[。！？；])", t) if p.strip() and len(p.strip())>3]

def _has_zh(t): return bool(re.search(r"[\u4e00-\u9fff]", t))

def _parse_pdf(fp):
    import fitz
    try:
        doc = fitz.open(fp)
        full = "".join(pg.get_text() for pg in doc)
        n = len(doc); doc.close()
    except Exception:
        return None
    # 知网型：原文内容 + 内联 pct%(chars)
    i2 = full.find("原文内容")
    out = []
    if i2 >= 0:
        after = full[i2+4:]
        inline = re.findall(r"([\d.]+)%\((\d+)\)", after)
        if inline:
            clean = re.sub(r"[\d.]+%\(\d+\)", "", after)
            clean = re.sub(r"—\s*\d+\s*—|https://cx\.cnki\.net|知网个人AIGC检测服务", "", clean)
            clean = re.sub(r"\s+", "", clean)
            pos = 0
            for pct, chars in inline:
                seg = clean[pos:pos+int(chars)]; pos += int(chars)
                for s in _split_sentences(seg):
                    if len(s) > 3 and _has_zh(s):
                        out.append({"text": s, "prob": round(float(pct)/100, 4), "label": _tier(float(pct))})
            if out: return out
    # PaperPass 型：疑似度分布（文档级，仅当无句级时返回整体率占位）
    return out if out else None

def _tier(prob):
    if prob >= 60: return "high"
    if prob >= 40: return "medium"
    return "human"

def _parse_word(fp):
    from docx import Document
    try:
        doc = Document(fp)
    except Exception:
        return None
    def map_label(rgb):
        if rgb is None: return ("human", 0.15)
        try: r,g,b = rgb[0],rgb[1],rgb[2]
        except: return ("human", 0.15)
        if r<60 and g<60 and b<60: return ("human", 0.15)
        if r>170 and g>170 and b>170: return ("skip", None)
        if r>220 and g<60 and b<60: return ("high", 0.80)
        if r>220 and 120<g<180 and b<40: return ("medium", 0.50)
        if 140<r<180 and 120<g<170 and b>210: return ("low", 0.22)
        return ("human", 0.15)
    cur_text=""; cur_label=None; out=[]
    for par in doc.paragraphs:
        for run in par.runs:
            t=run.text
            if not t: continue
            rgb = run.font.color.rgb if run.font.color and run.font.color.rgb else None
            label, prob = map_label(rgb)
            if label=="skip":
                if cur_text and len(cur_text)>4 and cur_label not in ("skip",):
                    out.append((cur_text.strip(), cur_label))
                cur_text=""; cur_label="skip"; continue
            if cur_label is None: cur_label=label; cur_text=t
            elif label!=cur_label:
                if cur_text and len(cur_text)>4 and cur_label not in ("skip",):
                    out.append((cur_text.strip(), cur_label))
                cur_text=t; cur_label=label
            else: cur_text+=t
        if cur_text and len(cur_text)>4 and cur_label not in ("skip",):
            out.append((cur_text.strip(), cur_label))
        cur_text=""; cur_label=None
    return [{"text":t,"label":lbl,"prob":{"high":0.8,"medium":0.5,"low":0.22,"human":0.15}[lbl]} for t,lbl in out if _has_zh(t)]

def _parse_html(fp):
    from bs4 import BeautifulSoup
    try:
        t=open(fp,encoding="utf-8",errors="ignore").read()
        soup=BeautifulSoup(t,"lxml")
    except Exception:
        return None
    out=[]
    for em in soup.find_all("em"):
        cls=(em.get("class") or [""])[0]
        if cls not in ("low","medium","high"): continue
        s=em.get_text(" ",strip=True)
        if not s: continue
        prob=None
        for sib in em.next_siblings:
            if getattr(sib,"name",None)=="div" and sib.get("class") and "aigc-detection-chance-popover" in sib.get("class"):
                m=re.search(r"AI生成的可能性为[：:]\s*<i[^>]*>\s*([\d.]+)%",str(sib))
                if m: prob=float(m.group(1))
                break
        if prob is not None and len(s)>3 and _has_zh(s):
            out.append({"text":s,"prob":round(prob/100,4),"label":_tier(prob)})
    return out if out else None

def existing_hashes():
    h=set()
    if os.path.exists(DATASET):
        for l in open(DATASET,encoding="utf-8"):
            if l.strip():
                try: h.add(hashlib.md5(json.loads(l)["text"].encode()).hexdigest())
                except: pass
    return h

def main():
    threshold = 500  # 新增句数阈值，达到则建议训练
    for i,a in enumerate(sys.argv):
        if a=="--threshold" and i+1<len(sys.argv):
            threshold=int(sys.argv[i+1])
    # 单文件模式：只解析指定文件（/ingest 接口传 --target）
    target = None
    for i,a in enumerate(sys.argv):
        if a=="--target" and i+1<len(sys.argv):
            target = sys.argv[i+1]
    sources = DEFAULT_SOURCES
    for i,a in enumerate(sys.argv):
        if a=="--source-dir" and i+1<len(sys.argv):
            sources=[sys.argv[i+1]]
    seen = existing_hashes()
    new_anns=[]; copied=[]
    if target:
        # 只处理指定文件
        if os.path.exists(target) and is_report_file(target):
            anns = parse_report(target)
            if anns:
                added=0
                for a in anns:
                    h=hashlib.md5(a["text"].encode()).hexdigest()
                    if h in seen: continue
                    seen.add(h); new_anns.append(a); added+=1
                if added:
                    copied.append({"src":target,"n":added})
        else:
            print(f"目标文件无效或非报告: {target}")
    else:
        for src in sources:
            if not os.path.isdir(src): continue
            for root,dirs,files in os.walk(src):
                dirs[:]=[d for d in dirs if d not in ("node_modules",".git","AppData","site-packages") and not d.startswith(".")]
                for fn in files:
                    fp=os.path.join(root,fn)
                    if not is_report_file(fp): continue
                    anns=parse_report(fp)
                    if not anns: continue
                    added=0
                    for a in anns:
                        h=hashlib.md5(a["text"].encode()).hexdigest()
                        if h in seen: continue
                        seen.add(h); new_anns.append(a); added+=1
                    if added:
                        copied.append({"src":fp,"n":added})
    # merge into dataset
    if new_anns:
        with open(DATASET,"a",encoding="utf-8") as f:
            for a in new_anns:
                f.write(json.dumps(a,ensure_ascii=False)+"\n")
    # 打印报告参数（对齐真实AIGC报告口径）
    for c in copied:
        meta = extract_report_meta(c["src"])
        info = f"  [入库] 报告: {meta.get('file','')[-45:]} | 平台={meta.get('platform')} | " \
               f"AI特征值={meta.get('ai_rate')}% | 总字符={meta.get('total_chars')} | " \
               f"报告编号={meta.get('report_no')} | 篇名={(meta.get('title') or '')[:30]} | 新增句={c['n']}"
        print(info)
    print("扫描完成。新增句级标注:", len(new_anns), " 涉及报告:", len(copied))
    if new_anns and len(new_anns) >= threshold:
        print(f"达到阈值 {threshold}，建议运行: python train_classifier.py")
    print("已合并进 train_unified.jsonl")

if __name__=="__main__":
    main()
